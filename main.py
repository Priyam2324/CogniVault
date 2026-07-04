from dotenv import load_dotenv
load_dotenv()
import os
import sys
import asyncio
import logging
import random
import sqlite3
import uuid
from datetime import datetime
from typing import List, Optional

import cognee
import pdfplumber
import pytesseract
import docx
from PIL import Image
from io import BytesIO

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import litellm

# -----------------------------------------------------------------------------
# ENV SETUP & DB
# -----------------------------------------------------------------------------
os.environ.setdefault("COGNEE_DIR", os.getenv("COGNEE_DIR", "./cognee_data"))
os.environ["ENABLE_BACKEND_ACCESS_CONTROL"] = "false"
os.environ["REQUIRE_AUTHENTICATION"] = "false"

if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

logger = logging.getLogger("cognivault")
logging.basicConfig(level=logging.INFO)

# -----------------------------------------------------------------------------
# LLM CONFIG — provider-agnostic, hot-swappable from .env alone
# -----------------------------------------------------------------------------
LLM_MODEL = os.getenv("LLM_MODEL")
if not LLM_MODEL:
    raise RuntimeError("LLM_MODEL environment variable is not set.")

LLM_API_KEY = os.getenv("LLM_API_KEY")
if not LLM_API_KEY:
    raise RuntimeError("LLM_API_KEY environment variable is not set.")

LLM_ENDPOINT = os.getenv("LLM_ENDPOINT")  # only set this for custom/self-hosted providers


async def call_llm(messages: list[dict], temperature: float = 0.7, timeout: float = 60.0) -> str:
    """Provider-agnostic chat completion. Swaps providers just by changing
    LLM_MODEL / LLM_API_KEY / LLM_ENDPOINT in .env - no code changes."""
    kwargs = {
        "model": LLM_MODEL,
        "messages": messages,
        "api_key": LLM_API_KEY,
        "temperature": temperature,
        "timeout": timeout,
    }
    if LLM_ENDPOINT:
        kwargs["api_base"] = LLM_ENDPOINT

    response = await litellm.acompletion(**kwargs)
    return response.choices[0].message.content

# -----------------------------------------------------------------------------
# ASSISTANT IDENTITY
# -----------------------------------------------------------------------------
CORE_IDENTITY = (
    "You are CogniVault, an AI assistant that reads documents - PDFs, photos, and "
    "Word docs - and remembers both them and past conversations across every chat "
    "a user has with you, not just the current one. When someone uploads a file, "
    "you extract its text (via OCR for images and scanned pages) and can answer "
    "questions about it later, in any conversation, without needing it re-uploaded."
)

FIRST_TURN_INSTRUCTION = (
    "This is the start of a new conversation. Briefly let the user know, in your own "
    "words and in one or two sentences, what you are and how you can help - keep it "
    "natural and conversational, not a canned feature list. Then address whatever "
    "they've actually said or uploaded."
)

CONTINUING_TURN_INSTRUCTION = (
    "This is a continuing conversation. Do NOT reintroduce yourself or re-explain "
    "what you are - the user already knows. Just answer."
)

# Guards against synthesis-time self-consistency slips: the model can extract two
# figures correctly and still botch the comparison between them when it writes the
# final sentence. Also pushes it to commit to an answer instead of hedging when the
# user has already told it what they care about.
SYNTHESIS_GUARDRAIL = (
    "Before finalizing any response that compares numbers, specs, or facts you cited "
    "earlier in the same response, re-check the comparison against the actual figures "
    "you just wrote down (e.g. if you said A=32 and B=48, verify which one is actually "
    "larger before stating a conclusion) - do not let the concluding sentence contradict "
    "numbers you already stated. If the user has told you what they care about or "
    "prioritize, use that to give a direct, decisive recommendation - do not fall back "
    "to 'it depends on your needs' when their stated priority already points to a clear "
    "answer based on your own analysis."
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
DB_PATH = "chat_history.db"


# -----------------------------------------------------------------------------
# SQLITE HELPER — WAL mode + busy_timeout + guaranteed close
# -----------------------------------------------------------------------------
# Plain `with sqlite3.connect(DB_PATH) as conn:` only wraps the transaction
# (commit/rollback on exit) - it does NOT close the connection. Every request
# was leaking a connection. Under a slow PDF/OCR upload (long-running request
# holding a connection open) overlapping with other requests (chat rename,
# get_chats polling, background memory writes), SQLite's default rollback
# journal (single writer, default 5s timeout) would throw "database is
# locked" - and because the inserts weren't wrapped in try/except, that
# exception aborted the whole request AFTER the LLM call had already run,
# so the message/chat update was silently dropped. WAL mode lets readers
# proceed while a write is in flight, a long busy_timeout makes concurrent
# writers wait instead of instantly erroring, and closing the connection
# every time stops the fd/connection leak entirely.
from contextlib import contextmanager


@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=30000")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def ensure_chat_exists(conn: sqlite3.Connection, chat_id: str, user_id: str) -> None:
    """Guards against orphaned messages. Previously, upload_file inserted
    messages against whatever chat_id the client sent without checking the
    chats table first. SQLite's declared FOREIGN KEY was never enforced
    (foreign_keys pragma defaults off), so if the client's "create chat" and
    "upload file" calls ever raced - or a chat_id was reused/mistyped - the
    messages got written but get_chats() (which only reads from `chats`)
    would never surface that chat. It looked exactly like the chat had been
    deleted, when really it had just never existed in the chats table."""
    row = conn.execute("SELECT 1 FROM chats WHERE id = ?", (chat_id,)).fetchone()
    if not row:
        chat_name = f"Chat-{datetime.now().strftime('%Y-%m-%d-%H:%M:%S')}"
        conn.execute(
            "INSERT INTO chats (id, user_id, name) VALUES (?, ?, ?)",
            (chat_id, user_id, chat_name),
        )


def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS chats (
                id TEXT PRIMARY KEY,
                user_id TEXT,
                name TEXT,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS messages (
                id TEXT PRIMARY KEY,
                chat_id TEXT,
                role TEXT,
                content TEXT,
                file_url TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(chat_id) REFERENCES chats(id)
            )
        """)
init_db()

# -----------------------------------------------------------------------------
# COGNEE LONG-TERM MEMORY
# -----------------------------------------------------------------------------
MEMORY_TOP_K = 8
MEMORY_MIN_SCORE = 0.68

async def recall_memories(
    query: str, user_id: str, top_k: int = MEMORY_TOP_K, min_score: float | None = MEMORY_MIN_SCORE
) -> list[str]:
    try:
        raw = await cognee.recall(query, session_id=user_id, top_k=top_k)
    except Exception:
        return []

    out = []
    for m in raw or []:
        if isinstance(m, dict):
            if min_score is not None and m.get("score") is not None and m.get("score") < min_score:
                continue
            text = m.get("answer") or m.get("text") or m.get("description")
        else:
            text = str(m)
        if text and len(text.strip()) >= 5:
            out.append(text.strip())
    return out[:top_k]

async def store_memory(user_id: str, data: str) -> None:
    try:
        await cognee.remember(data, session_id=user_id)
    except Exception:
        logger.exception("cognee.remember failed")

# -----------------------------------------------------------------------------
# DOCUMENT INGESTION - tuned to minimize token spend
# -----------------------------------------------------------------------------
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "6000"))
CHUNK_STORE_DELAY_SECONDS = float(os.getenv("CHUNK_STORE_DELAY_SECONDS", "1.0"))


async def store_document_chunks(user_id: str, chunks: List[str], filename: str) -> None:
    for chunk in chunks:
        await store_memory(user_id, f"Document '{filename}': {chunk}")
        await asyncio.sleep(CHUNK_STORE_DELAY_SECONDS)


# -----------------------------------------------------------------------------
# FAIR MULTI-FILE CONTEXT BUDGETING (for the immediate post-upload reply)
# -----------------------------------------------------------------------------
UPLOAD_CONTEXT_CHAR_BUDGET = int(os.getenv("UPLOAD_CONTEXT_CHAR_BUDGET", "40000"))


def allocate_char_budgets(lengths: list[int], total_budget: int) -> list[int]:
    n = len(lengths)
    if n == 0:
        return []
    budgets = [0] * n
    remaining_budget = total_budget
    remaining_files = n
    for idx in sorted(range(n), key=lambda i: lengths[i]):
        fair_share = remaining_budget // remaining_files
        take = min(lengths[idx], fair_share)
        budgets[idx] = take
        remaining_budget -= take
        remaining_files -= 1
    return budgets

# Helper to run heavy synchronous file parsing off the main event loop
def extract_text_sync(file_path: str, ext: str, contents: bytes) -> str:
    extracted_text = ""
    try:
        if ext == "pdf":
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text)

                    for t_idx, table in enumerate(page.extract_tables(), start=1):
                        rows = []
                        for row in table:
                            cells = [(c or "").strip() for c in row]
                            if any(cells):
                                rows.append(" | ".join(cells))
                        if rows:
                            parts.append(f"[Table {t_idx}, page {page_num}]\n" + "\n".join(rows))
            extracted_text = "\n\n".join(parts)

        elif ext in ["png", "jpg", "jpeg"]:
            img = Image.open(BytesIO(contents))
            extracted_text = pytesseract.image_to_string(img)

        elif ext in ["docx", "doc"]:
            doc = docx.Document(BytesIO(contents))
            parts = [para.text for para in doc.paragraphs if para.text]

            for t_idx, table in enumerate(doc.tables, start=1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append(f"[Table {t_idx}]\n" + "\n".join(rows))

            extracted_text = "\n\n".join(parts)
    except Exception as e:
        extracted_text = f"[Error extracting text from {ext} file: {str(e)}]"
    return extracted_text


# -----------------------------------------------------------------------------
# FASTAPI APP & ENDPOINTS
# -----------------------------------------------------------------------------
app = FastAPI(title="CogniVault Document Agent")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

class ChatCreate(BaseModel):
    user_id: str
    name: Optional[str] = None

class ChatRename(BaseModel):
    name: str

class ChatMessage(BaseModel):
    message: str
    chat_id: str
    user_id: str

@app.post("/api/chats")
def create_chat(payload: ChatCreate):
    chat_id = str(uuid.uuid4())

    with get_db() as conn:
        chat_name = payload.name
        if not chat_name or chat_name == "New Chat":
            current_time = datetime.now().strftime("%Y-%m-%d-%H:%M:%S")
            chat_name = f"Chat-{current_time}"

        conn.execute("INSERT INTO chats (id, user_id, name) VALUES (?, ?, ?)", (chat_id, payload.user_id, chat_name))

    return {"id": chat_id, "name": chat_name}

@app.get("/api/chats")
def get_chats(user_id: str):
    with get_db() as conn:
        conn.row_factory = sqlite3.Row
        chats = conn.execute("SELECT * FROM chats WHERE user_id = ? ORDER BY updated_at DESC", (user_id,)).fetchall()
    return [dict(c) for c in chats]

@app.put("/api/chats/{chat_id}")
def rename_chat(chat_id: str, payload: ChatRename):
    with get_db() as conn:
        conn.execute("UPDATE chats SET name = ? WHERE id = ?", (payload.name, chat_id))
    return {"status": "success"}

@app.delete("/api/chats/{chat_id}")
def delete_chat(chat_id: str):
    with get_db() as conn:
        conn.execute("DELETE FROM messages WHERE chat_id = ?", (chat_id,))
        conn.execute("DELETE FROM chats WHERE id = ?", (chat_id,))
    return {"status": "deleted"}

@app.get("/api/chats/{chat_id}/messages")
def get_messages(chat_id: str):
    with get_db() as conn:
        conn.row_factory = sqlite3.Row
        msgs = conn.execute("SELECT * FROM messages WHERE chat_id = ? ORDER BY created_at ASC", (chat_id,)).fetchall()
    return [dict(m) for m in msgs]

@app.post("/api/upload")
async def upload_file(
    chat_id: str = Form(...),
    user_id: str = Form(...),
    prompt: str = Form(""),
    files: List[UploadFile] = File(...)
):
    # Make sure the chat this upload is attributed to actually exists before
    # anything gets written against it - see ensure_chat_exists() above.
    try:
        with get_db() as conn:
            ensure_chat_exists(conn, chat_id, user_id)
            existing_count = conn.execute(
                "SELECT COUNT(*) FROM messages WHERE chat_id = ?", (chat_id,)
            ).fetchone()[0]
    except Exception:
        logger.exception("Failed to verify/create chat %s before upload", chat_id)
        raise HTTPException(status_code=500, detail="Could not initialize chat for upload.")

    first_turn = existing_count == 0

    file_texts = []
    file_urls = []
    file_names = []

    for file in files:
        file_id = str(uuid.uuid4())
        ext = file.filename.split('.')[-1].lower()
        file_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")

        contents = await file.read()
        with open(file_path, "wb") as f:
            f.write(contents)

        extracted_text = await asyncio.to_thread(extract_text_sync, file_path, ext, contents)

        if extracted_text.strip():
            text_clean = extracted_text.replace("\n", " ").strip()
            chunks = [text_clean[i:i+CHUNK_SIZE] for i in range(0, len(text_clean), CHUNK_SIZE)]
            asyncio.create_task(store_document_chunks(user_id, chunks, file.filename))

        file_texts.append((file.filename, extracted_text))
        file_urls.append((file.filename, f"/uploads/{file_id}.{ext}"))
        file_names.append(file.filename)

    try:
        with get_db() as conn:
            for fname, furl in file_urls:
                conn.execute(
                    "INSERT INTO messages (id, chat_id, role, content, file_url) VALUES (?, ?, ?, ?, ?)",
                    (str(uuid.uuid4()), chat_id, "user", f"Uploaded file: {fname}", furl)
                )
    except Exception:
        logger.exception("Failed to save uploaded-file messages for chat %s", chat_id)
        raise HTTPException(status_code=500, detail="Files were processed but could not be saved to chat history.")

    names_str = ", ".join(file_names)

    if prompt.strip():
        llm_prompt = prompt
        try:
            with get_db() as conn:
                conn.execute("INSERT INTO messages (id, chat_id, role, content) VALUES (?, ?, ?, ?)",
                             (str(uuid.uuid4()), chat_id, "user", prompt))
        except Exception:
            logger.exception("Failed to save user prompt for chat %s", chat_id)
            raise HTTPException(status_code=500, detail="Could not save your message to chat history.")
    else:
        llm_prompt = f"I have uploaded files: {names_str}. Please provide a brief summary of their contents."

    lengths = [len(text) for _, text in file_texts]
    budgets = allocate_char_budgets(lengths, UPLOAD_CONTEXT_CHAR_BUDGET)

    extracted_blocks = []
    for (fname, text), budget in zip(file_texts, budgets):
        truncated = len(text) > budget
        block_text = text[:budget]
        if truncated:
            block_text += (
                f"\n\n[NOTE: '{fname}' is long and was truncated to fit context - "
                f"only the first {budget:,} of {len(text):,} extracted characters "
                f"are shown above. Say so if the answer might depend on a part "
                f"that was cut off.]"
            )
        extracted_blocks.append(f"--- Content of {fname} ---\n{block_text}")

    all_extracted_text = "\n\n".join(extracted_blocks)

    turn_instruction = FIRST_TURN_INSTRUCTION if first_turn else CONTINUING_TURN_INSTRUCTION
    sys_prompt = (
        f"{CORE_IDENTITY}\n\n"
        f"{turn_instruction}\n\n"
        f"{SYNTHESIS_GUARDRAIL}\n\n"
        f"The user just uploaded: {names_str}. Answer their prompt based strictly on "
        f"the extracted text below - don't invent details that aren't in it, and say "
        f"so plainly if the extracted text doesn't cover what they asked.\n\n"
        f"EXTRACTED TEXT:\n---\n{all_extracted_text}\n---"
    )

    messages = [{"role": "system", "content": sys_prompt}, {"role": "user", "content": llm_prompt}]

    try:
        response_text = await call_llm(messages, timeout=90.0)
    except Exception as e:
        response_text = f"Files processed, but AI generation failed: {str(e)}"

    try:
        with get_db() as conn:
            conn.execute("INSERT INTO messages (id, chat_id, role, content) VALUES (?, ?, ?, ?)",
                        (str(uuid.uuid4()), chat_id, "assistant", response_text))
            conn.execute("UPDATE chats SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (chat_id,))
    except Exception:
        logger.exception("Failed to save assistant response for chat %s", chat_id)
        # The reply was already generated - still return it to the user even
        # though persistence failed, and log loudly so it's not "random".
        return {"response": response_text, "files_processed": len(file_urls), "warning": "Response could not be saved to history."}

    asyncio.create_task(store_memory(
        user_id,
        f"User uploaded document(s) {names_str} and asked: '{llm_prompt}'. Answer given: {response_text}"
    ))

    return {"response": response_text, "files_processed": len(file_urls)}

@app.post("/api/chat")
async def chat(payload: ChatMessage):
    try:
        with get_db() as conn:
            ensure_chat_exists(conn, payload.chat_id, payload.user_id)
            conn.row_factory = sqlite3.Row
            history = conn.execute("SELECT role, content FROM messages WHERE chat_id = ? ORDER BY created_at DESC LIMIT 12", (payload.chat_id,)).fetchall()
            history = list(reversed(history))
    except Exception:
        logger.exception("Failed to load history for chat %s", payload.chat_id)
        raise HTTPException(status_code=500, detail="Could not load chat history.")

    history_msgs = [{"role": row["role"], "content": row["content"]} for row in history]
    first_turn = len(history_msgs) == 0
    turn_instruction = FIRST_TURN_INSTRUCTION if first_turn else CONTINUING_TURN_INSTRUCTION

    recent_context = " ".join(m["content"] for m in history_msgs[-4:])
    recall_query = f"{recent_context} {payload.message}".strip()
    memories = await recall_memories(recall_query, payload.user_id)

    if memories:
        memory_block = "\n".join(f"- {m}" for m in memories)
        sys_prompt = (
            f"{CORE_IDENTITY}\n\n"
            f"{turn_instruction}\n\n"
            f"{SYNTHESIS_GUARDRAIL}\n\n"
            "Answer the user's question directly, naturally, and concisely - the way a "
            "knowledgeable person would, not like a system reporting on its own internals.\n\n"
            "Below is some context remembered from earlier conversations with this user, "
            "including anything relevant from documents they've uploaded. Use a piece of it "
            "ONLY if it directly helps answer the current question. If none of it is relevant, "
            "ignore all of it completely and answer from your own knowledge instead - do not "
            "mention, list, or summarize any of these facts just because they exist. Bringing "
            "up unrelated personal details or documents when they weren't asked about is not "
            "helpful, it's noise - never do this as a consolation when you don't know the "
            "actual answer.\n\n"
            "Never mention that you 'retrieved', 'recall', or 'remember' this context, never "
            "say things like 'from our previous conversation' or 'I don't have X, but I can "
            "tell you about Y' - just answer the question that was actually asked.\n\n"
            "Also check the conversation history below: if something has already been "
            "established earlier in THIS conversation, don't restate it again on follow-up "
            "questions unless the user actually asks for it again - just answer the new "
            "question directly, assuming that context is already shared.\n\n"
            f"REMEMBERED CONTEXT:\n{memory_block}"
        )
    else:
        sys_prompt = (
            f"{CORE_IDENTITY}\n\n"
            f"{turn_instruction}\n\n"
            f"{SYNTHESIS_GUARDRAIL}\n\n"
            "Answer the user's question directly, naturally, and concisely, using your own "
            "knowledge. Do not mention memory, recall, or previous conversations."
        )

    messages = [{"role": "system", "content": sys_prompt}] + history_msgs + [{"role": "user", "content": payload.message}]

    try:
        response_text = await call_llm(messages, timeout=60.0)
    except Exception as e:
        response_text = f"AI generation failed: {str(e)}"

    try:
        with get_db() as conn:
            conn.execute("INSERT INTO messages (id, chat_id, role, content) VALUES (?, ?, ?, ?)", (str(uuid.uuid4()), payload.chat_id, "user", payload.message))
            conn.execute("INSERT INTO messages (id, chat_id, role, content) VALUES (?, ?, ?, ?)", (str(uuid.uuid4()), payload.chat_id, "assistant", response_text))
            conn.execute("UPDATE chats SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (payload.chat_id,))
    except Exception:
        logger.exception("Failed to save chat turn for chat %s", payload.chat_id)
        return {"response": response_text, "warning": "Response could not be saved to history."}

    asyncio.create_task(store_memory(payload.user_id, f"User: {payload.message} | Assistant: {response_text}"))
    return {"response": response_text}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)